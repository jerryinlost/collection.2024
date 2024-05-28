import os
import csv
from tabulate import tabulate
from docx import Document

def get_largest_model_per_directory(root_dir):
    """
    Walk through the directory and find the largest model file in each subdirectory.
    
    :param root_dir: The root directory containing the models.
    :return: List of lists containing structured path and size for each largest model found.
    """
    # Define the model file extensions to look for
    model_extensions = {'.pt', '.pth', '.bin', '.onnx'}
    
    # Dictionary to store the largest model file in each directory
    largest_models = {}

    # Walk through the directory
    for dirpath, dirnames, filenames in os.walk(root_dir):
        largest_file = None
        largest_file_size = 0
        
        for filename in filenames:
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension in model_extensions:
                file_path = os.path.join(dirpath, filename)
                file_size = os.path.getsize(file_path)  # Size in bytes
                
                # Check if this is the largest file found in this directory
                if file_size > largest_file_size:
                    largest_file = filename
                    largest_file_size = file_size

        if largest_file:
            relative_path = os.path.relpath(dirpath, root_dir)
            structured_path = relative_path.replace(f"{os.path.basename(root_dir)}/", "")
            largest_models[structured_path] = largest_file_size

    # Prepare the table data
    table_data = [[path, f"{size / (1024 * 1024):.2f} MB"] for path, size in largest_models.items()]

    return table_data

def export_to_docx(data, docx_file):
    """
    Export the data to a DOCX file.
    
    :param data: List of lists containing structured path and size.
    :param docx_file: Path to the DOCX file to be created.
    """
    # Create a new Document
    doc = Document()
    
    # Add a title to the document
    doc.add_heading('Model Sizes', level=1)
    
    # Add a table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model Name (Structured Path)'
    hdr_cells[1].text = 'Size'
    
    # Add the data rows
    for row in data:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
    
    # Save the document
    doc.save(docx_file)
    print(f"Table data exported to {docx_file}")

def main():
    # Specify the root directory containing the models
    root_directory = './models'
    
    # Get the largest models data
    table_data = get_largest_model_per_directory(root_directory)
    
    # Print the table
    table_headers = ["Model Name (Structured Path)", "Size"]
    print(tabulate(table_data, headers=table_headers, tablefmt="grid"))
    
    # Export to DOCX
    docx_file = 'largest_models.docx'
    export_to_docx(table_data, docx_file)

if __name__ == "__main__":
    main()